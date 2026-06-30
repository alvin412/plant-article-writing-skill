#!/usr/bin/env python3
"""Build the three final Word deliverables for plant article workflows."""

from __future__ import annotations

import argparse
import json
import re
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_NAMES = {
    "manuscript": "01_plant_article_bilingual.docx",
    "audit": "02_claim_citation_audit_bilingual.docx",
    "review": "03_peer_review_and_response_bilingual.docx",
}

PLACEHOLDER_PATTERNS = (
    r"\[\[.+?\]\]",
    r"\bTBD\b",
    r"\bTODO\b",
    r"AUTHOR_INPUT_NEEDED",
    r"EVIDENCE_NEEDED",
)


def template_payload() -> dict[str, Any]:
    return {
        "project": {
            "title_en": "[[English title]]",
            "title_zh": "[[中文标题]]",
            "article_type": "Review or Original Research",
            "target_journal": "[[Target journal or Not selected]]",
            "authors": "[[Authors]]",
            "version": "v1.0",
            "date": date.today().isoformat(),
        },
        "manuscript": {
            "english": [
                {"heading": "Title", "level": 1, "paragraphs": ["[[English title]]"]},
                {"heading": "Abstract", "level": 1, "paragraphs": ["[[English abstract]]"]},
                {"heading": "Keywords", "level": 1, "paragraphs": ["[[English keywords]]"]},
                {"heading": "Introduction", "level": 1, "paragraphs": ["[[English introduction]]"]},
                {"heading": "References", "level": 1, "paragraphs": ["[[Verified references]]"]},
            ],
            "chinese": [
                {"heading": "题目", "level": 1, "paragraphs": ["[[中文题目]]"]},
                {"heading": "摘要", "level": 1, "paragraphs": ["[[中文摘要]]"]},
                {"heading": "关键词", "level": 1, "paragraphs": ["[[中文关键词]]"]},
                {"heading": "引言", "level": 1, "paragraphs": ["[[中文引言]]"]},
                {"heading": "参考文献", "level": 1, "paragraphs": ["[[同一组已核验参考文献]]"]},
            ],
        },
        "audit": {
            "summary_en": "[[English audit summary]]",
            "summary_zh": "[[中文核查摘要]]",
            "counts": {"claims": 0, "pass": 0, "warn": 0, "block": 0},
            "data_consistency_en": "[[Data consistency summary or Not applicable]]",
            "data_consistency_zh": "[[数据一致性摘要或不适用]]",
            "claim_citation_matrix": [
                {
                    "claim_id": "C001",
                    "location": "[[Section and paragraph]]",
                    "claim_en": "[[English manuscript claim]]",
                    "claim_zh": "[[中文论点]]",
                    "citation": "[[Citation key]]",
                    "support_grade": "[[strong/partial/background/contradictory]]",
                    "boundary": "[[Evidence boundary]]",
                    "decision": "[[PASS/WARN/BLOCK]]",
                }
            ],
            "sources": [
                {
                    "citation_key": "[[Citation key]]",
                    "title": "[[Paper title]]",
                    "doi": "[[DOI]]",
                    "zotero_key": "[[Zotero item key]]",
                    "reading_status": "[[full-text/partial-reading]]",
                    "uses": [
                        {
                            "claim_id": "C001",
                            "source_anchor": "[[Page/section/paragraph/figure/table]]",
                            "original_excerpt": "[[Short supporting excerpt]]",
                            "translation_zh": "[[学术中文翻译]]",
                            "rationale_en": "[[English support explanation]]",
                            "rationale_zh": "[[中文支持说明]]",
                            "boundary": "[[Scope boundary]]",
                        }
                    ],
                }
            ],
            "unresolved_en": ["[[Unresolved risk or None]]"],
            "unresolved_zh": ["[[未解决风险或无]]"],
        },
        "peer_review": {
            "editorial_synthesis_en": "[[English editorial synthesis]]",
            "editorial_synthesis_zh": "[[中文编辑综合意见]]",
            "final_posture_en": "[[Final re-review posture]]",
            "final_posture_zh": "[[最终复审结论]]",
            "reviewers": [
                {
                    "name_en": "Reviewer 1 - Plant Domain",
                    "name_zh": "审稿人1——植物领域",
                    "overall_en": "[[English overall assessment]]",
                    "overall_zh": "[[中文总体评价]]",
                    "comments": [
                        {
                            "id": "R1.1",
                            "severity": "Major",
                            "location": "[[Location]]",
                            "comment_en": "[[English reviewer comment]]",
                            "comment_zh": "[[中文审稿意见]]",
                            "response_en": "[[English point-by-point response]]",
                            "response_zh": "[[中文逐点回复]]",
                            "change_en": "[[English exact manuscript change]]",
                            "change_zh": "[[中文具体修改]]",
                            "verification_en": "[[English re-review verification]]",
                            "verification_zh": "[[中文复审核验]]",
                            "status": "open",
                        }
                    ],
                }
            ],
            "residual_en": ["[[Residual risk or None]]"],
            "residual_zh": ["[[剩余风险或无]]"],
        },
    }


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def set_shading(cell, fill: str = "D9EAF7") -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title = doc.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 78, 121)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for name, size, color in (
        ("Heading 1", 15, RGBColor(31, 78, 121)),
        ("Heading 2", 12.5, RGBColor(55, 55, 55)),
        ("Heading 3", 11, RGBColor(80, 80, 80)),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.keep_with_next = True

    add_page_number(section.footer.paragraphs[0])


def add_text(doc: Document, value: Any, label: str | None = None) -> None:
    if value is None:
        return
    values = value if isinstance(value, list) else [value]
    for item in values:
        paragraph = doc.add_paragraph()
        if label:
            paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(str(item))


def add_key_value_table(doc: Document, rows: list[tuple[str, Any]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[0].paragraphs[0].add_run(str(label)).bold = True
        cells[1].paragraphs[0].add_run(str(value or ""))
        set_shading(cells[0])


def add_table(doc: Document, rows: list[dict[str, Any]], columns: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    repeat_header(table.rows[0])
    for index, (_, label) in enumerate(columns):
        cell = table.rows[0].cells[index]
        cell.text = label
        set_shading(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, (key, _) in enumerate(columns):
            cells[index].text = str(row.get(key, ""))
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def metadata(payload: dict[str, Any]) -> list[tuple[str, Any]]:
    project = payload.get("project", {})
    return [
        ("English title", project.get("title_en", "")),
        ("Chinese title", project.get("title_zh", "")),
        ("Article type", project.get("article_type", "")),
        ("Target journal", project.get("target_journal", "")),
        ("Authors", project.get("authors", "")),
        ("Version", project.get("version", "")),
        ("Date", project.get("date", "")),
    ]


def add_title_page(doc: Document, title: str, payload: dict[str, Any]) -> None:
    paragraph = doc.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(title)
    add_key_value_table(doc, metadata(payload))
    doc.add_page_break()


def add_sections(doc: Document, sections: list[dict[str, Any]], offset: int = 0) -> None:
    for section in sections:
        level = min(max(int(section.get("level", 1)) + offset, 1), 3)
        doc.add_heading(str(section.get("heading", "Untitled")), level=level)
        for paragraph in section.get("paragraphs", []):
            add_text(doc, paragraph)


def build_manuscript(payload: dict[str, Any], path: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_title_page(doc, "Bilingual Plant-Science Article / 植物科学双语正文", payload)
    doc.add_heading("English Manuscript", level=1)
    add_sections(doc, payload.get("manuscript", {}).get("english", []), offset=1)
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Academic Chinese Manuscript / 学术中文正文", level=1)
    add_sections(doc, payload.get("manuscript", {}).get("chinese", []), offset=1)
    doc.save(path)


def build_audit(payload: dict[str, Any], path: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_title_page(doc, "Claim and Citation Audit / 论点与引文核查", payload)
    audit = payload.get("audit", {})
    doc.add_heading("Audit Summary", level=1)
    add_text(doc, audit.get("summary_en", ""))
    doc.add_heading("核查摘要", level=1)
    add_text(doc, audit.get("summary_zh", ""))
    doc.add_heading("Counts / 统计", level=1)
    add_key_value_table(doc, list(audit.get("counts", {}).items()))
    doc.add_heading("Data Consistency / 数据一致性", level=1)
    add_text(doc, audit.get("data_consistency_en", ""), "English")
    add_text(doc, audit.get("data_consistency_zh", ""), "中文")
    doc.add_heading("Claim-Citation Matrix / 论点—引文矩阵", level=1)
    add_table(
        doc,
        audit.get("claim_citation_matrix", []),
        [
            ("claim_id", "ID"),
            ("location", "Location"),
            ("claim_en", "Claim"),
            ("citation", "Citation"),
            ("support_grade", "Support"),
            ("boundary", "Boundary"),
            ("decision", "Decision"),
        ],
    )
    for index, source in enumerate(audit.get("sources", []), start=1):
        doc.add_page_break()
        doc.add_heading(
            f"Source {index}: {source.get('citation_key', '')} — {source.get('title', '')}",
            level=1,
        )
        add_key_value_table(
            doc,
            [
                ("DOI", source.get("doi", "")),
                ("Zotero item key", source.get("zotero_key", "")),
                ("Reading status", source.get("reading_status", "")),
            ],
        )
        for use in source.get("uses", []):
            doc.add_heading(f"Citation use {use.get('claim_id', '')}", level=2)
            for label, key in (
                ("Source anchor", "source_anchor"),
                ("Original excerpt", "original_excerpt"),
                ("学术中文翻译", "translation_zh"),
                ("Support rationale", "rationale_en"),
                ("支持说明", "rationale_zh"),
                ("Boundary", "boundary"),
            ):
                add_text(doc, use.get(key, ""), label)
    doc.add_heading("Unresolved Risks", level=1)
    add_text(doc, audit.get("unresolved_en", []))
    doc.add_heading("未解决风险", level=1)
    add_text(doc, audit.get("unresolved_zh", []))
    doc.save(path)


def build_review(payload: dict[str, Any], path: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_title_page(doc, "Peer Review and Point-by-Point Response / 审稿意见与逐点回复", payload)
    review = payload.get("peer_review", {})
    doc.add_heading("Editorial Synthesis", level=1)
    add_text(doc, review.get("editorial_synthesis_en", ""))
    doc.add_heading("编辑综合意见", level=1)
    add_text(doc, review.get("editorial_synthesis_zh", ""))
    for reviewer in review.get("reviewers", []):
        doc.add_page_break()
        doc.add_heading(
            f"{reviewer.get('name_en', 'Reviewer')} / {reviewer.get('name_zh', '审稿人')}",
            level=1,
        )
        add_text(doc, reviewer.get("overall_en", ""), "Overall assessment")
        add_text(doc, reviewer.get("overall_zh", ""), "总体评价")
        for comment in reviewer.get("comments", []):
            doc.add_heading(
                f"{comment.get('id', 'Comment')} — {comment.get('severity', '')}",
                level=2,
            )
            for label, key in (
                ("Location", "location"),
                ("Reviewer comment", "comment_en"),
                ("审稿意见", "comment_zh"),
                ("Author response", "response_en"),
                ("作者回复", "response_zh"),
                ("Exact manuscript change", "change_en"),
                ("具体修改", "change_zh"),
                ("Re-review verification", "verification_en"),
                ("复审核验", "verification_zh"),
                ("Status", "status"),
            ):
                add_text(doc, comment.get(key, ""), label)
    doc.add_heading("Final Re-Review Posture", level=1)
    add_text(doc, review.get("final_posture_en", ""))
    doc.add_heading("最终复审结论", level=1)
    add_text(doc, review.get("final_posture_zh", ""))
    doc.add_heading("Residual Risks", level=1)
    add_text(doc, review.get("residual_en", []))
    doc.add_heading("剩余风险", level=1)
    add_text(doc, review.get("residual_zh", []))
    doc.save(path)


def collect_strings(value: Any) -> list[str]:
    if isinstance(value, dict):
        result: list[str] = []
        for child in value.values():
            result.extend(collect_strings(child))
        return result
    if isinstance(value, list):
        result = []
        for child in value:
            result.extend(collect_strings(child))
        return result
    return [str(value)] if value is not None else []


def validate_docx(path: Path) -> dict[str, Any]:
    doc = Document(path)
    titles = [p.text.strip() for p in doc.paragraphs if p.style and p.style.name == "Title"]
    levels: list[int] = []
    errors: list[str] = []
    previous = 0
    for paragraph in doc.paragraphs:
        match = re.fullmatch(r"Heading ([1-3])", paragraph.style.name if paragraph.style else "")
        if not match:
            continue
        level = int(match.group(1))
        levels.append(level)
        if previous and level > previous + 1:
            errors.append(f"Heading jump {previous}->{level}: {paragraph.text[:80]}")
        previous = level
    return {
        "path": str(path.resolve()),
        "reopens": True,
        "nonempty_title": bool(titles and titles[0]),
        "heading_count": len(levels),
        "heading_hierarchy_errors": errors,
        "table_count": len(doc.tables),
        "paragraph_count": len(doc.paragraphs),
    }


def build_package(payload: dict[str, Any], project_dir: Path) -> dict[str, Any]:
    deliverables = project_dir / "deliverables"
    intermediate = project_dir / "intermediate_files"
    validation_dir = intermediate / "validation"
    deliverables.mkdir(parents=True, exist_ok=True)
    validation_dir.mkdir(parents=True, exist_ok=True)

    paths = {key: deliverables / filename for key, filename in OUTPUT_NAMES.items()}
    build_manuscript(payload, paths["manuscript"])
    build_audit(payload, paths["audit"])
    build_review(payload, paths["review"])

    all_text = "\n".join(collect_strings(payload))
    placeholders = sorted(
        {
            match.group(0)
            for pattern in PLACEHOLDER_PATTERNS
            for match in re.finditer(pattern, all_text, flags=re.IGNORECASE)
        }
    )
    report = {
        "generated_on": date.today().isoformat(),
        "files": [validate_docx(path) for path in paths.values()],
        "placeholder_warnings": placeholders,
        "render_status": "not-run",
        "release_status": "BLOCK_PLACEHOLDERS" if placeholders else "STRUCTURAL_PASS_RENDER_REQUIRED",
    }
    report_path = validation_dir / "word_package_validation.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def main() -> int:
    parser = argparse.ArgumentParser(description="Build the three plant-article Word deliverables.")
    parser.add_argument("input_json", nargs="?", type=Path)
    parser.add_argument("--project-dir", type=Path, default=Path("article_project"))
    parser.add_argument("--init-template", action="store_true")
    args = parser.parse_args()

    if args.init_template:
        destination = args.project_dir / "intermediate_files" / "word_inputs" / "project.json"
        destination.parent.mkdir(parents=True, exist_ok=True)
        destination.write_text(
            json.dumps(template_payload(), ensure_ascii=False, indent=2), encoding="utf-8"
        )
        print(f"Template written: {destination.resolve()}")
        return 0

    if not args.input_json:
        parser.error("input_json is required unless --init-template is used")

    payload = json.loads(args.input_json.read_text(encoding="utf-8-sig"))
    print(json.dumps(build_package(payload, args.project_dir), ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
