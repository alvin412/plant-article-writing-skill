#!/usr/bin/env python3
"""Build the three required plant-article-writing Word deliverables."""

from __future__ import annotations

import argparse
import json
import re
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_NAMES = {
    "manuscript": "01_plant_manuscript_bilingual.docx",
    "review": "02_peer_review_and_response.docx",
    "literature": "03_literature_audit_and_close_reading.docx",
}

PLACEHOLDER_PATTERNS = (
    r"\[\[.+?\]\]",
    r"\[Evidence needed",
    r"\bTBD\b",
    r"\bTODO\b",
    r"AUTHOR_INPUT_NEEDED",
)


def template_payload() -> dict[str, Any]:
    return {
        "project": {
            "title_en": "[[English project title]]",
            "title_zh": "[[Chinese project title]]",
            "manuscript_type": "Review",
            "target_journal": "[[Target journal]]",
            "authors": "[[Authors]]",
            "version": "v1.0",
            "date": date.today().isoformat(),
        },
        "manuscript": {
            "english": [
                {
                    "heading": "Abstract",
                    "level": 1,
                    "paragraphs": ["[[English abstract]]"],
                },
                {
                    "heading": "Introduction",
                    "level": 1,
                    "paragraphs": ["[[English introduction]]"],
                },
                {
                    "heading": "References",
                    "level": 1,
                    "paragraphs": ["[[Formatted reference list]]"],
                },
            ],
            "chinese": [
                {
                    "heading": "摘要",
                    "level": 1,
                    "paragraphs": ["[[Chinese abstract]]"],
                },
                {
                    "heading": "引言",
                    "level": 1,
                    "paragraphs": ["[[Chinese introduction]]"],
                },
                {
                    "heading": "参考文献",
                    "level": 1,
                    "paragraphs": ["[[Same verified reference list]]"],
                },
            ],
        },
        "review": {
            "editorial_synthesis": "[[Editorial synthesis]]",
            "final_decision": "[[Final re-review posture]]",
            "residual_risks": ["[[Residual risk or None]]"],
            "reviewers": [
                {
                    "name": "Reviewer 1 - Plant Domain",
                    "overall_assessment": "[[Overall assessment]]",
                    "comments": [
                        {
                            "id": "R1.1",
                            "severity": "Major",
                            "comment": "[[Preserved reviewer comment]]",
                            "response": "[[Point-by-point author response]]",
                            "change": "[[Exact manuscript change]]",
                            "location": "[[Section/paragraph/line/figure/table]]",
                            "verification": "[[Re-review verification]]",
                            "status": "open",
                        }
                    ],
                }
            ],
        },
        "literature": {
            "search_record": {
                "databases": "[[Databases]]",
                "search_dates": "[[Search dates]]",
                "search_strings": "[[Exact search strings]]",
                "inclusion_exclusion": "[[Criteria]]",
                "screening_counts": "[[Identification and screening counts]]",
            },
            "audit_summary": {
                "total_references": 0,
                "verified": 0,
                "warnings": 0,
                "blocked": 0,
            },
            "claim_citation_matrix": [
                {
                    "claim_id": "C001",
                    "manuscript_claim": "[[Manuscript claim]]",
                    "citation": "[[Citation key]]",
                    "support_grade": "[[strong/partial/background/contradictory]]",
                    "status": "[[PASS/WARN/BLOCK]]",
                    "boundary": "[[Species/tissue/treatment/evidence boundary]]",
                }
            ],
            "papers": [
                {
                    "citation_key": "[[Citation key]]",
                    "title": "[[Paper title]]",
                    "authors": "[[Authors]]",
                    "journal": "[[Journal]]",
                    "year": "[[Year]]",
                    "doi": "[[DOI]]",
                    "pmid": "[[PMID/PMCID]]",
                    "url": "[[URL]]",
                    "access_status": "short-excerpt",
                    "publication_status": "[[Published/corrected/retracted]]",
                    "plant_context": "[[Species/tissue/treatment/trait]]",
                    "close_reading": [
                        {
                            "section": "Abstract",
                            "source_anchors": "[[Page/section/block]]",
                            "note_en": "[[English analytical note]]",
                            "note_zh": "[[Academic Chinese note]]",
                            "key_evidence": "[[Evidence]]",
                            "figures_tables": "[[Figures/tables]]",
                            "limitations": "[[Limitations]]",
                            "manuscript_uses": "[[Claim IDs or sections]]",
                        }
                    ],
                    "citation_uses": [
                        {
                            "claim_id": "C001",
                            "manuscript_sentence": "[[Cited manuscript sentence]]",
                            "manuscript_location": "[[Section/paragraph]]",
                            "source_anchor": "[[Page/section/paragraph/figure]]",
                            "original_excerpt": "[[Short exact supporting excerpt]]",
                            "translation_zh": "[[Academic Chinese translation]]",
                            "support_grade": "[[Support grade]]",
                            "boundary": "[[Support boundary]]",
                        }
                    ],
                }
            ],
            "unresolved_evidence_gaps": ["[[Gap or None]]"],
        },
    }


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_in: list[float], indent_dxa: int = 120) -> None:
    widths_dxa = [round(width * 1440) for width in widths_in]
    total_dxa = sum(widths_dxa)
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_dxa in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width_dxa))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width_dxa in zip(row.cells, widths_dxa):
            cell.width = Inches(width_dxa / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, margin in (
                ("top", 80),
                ("bottom", 80),
                ("start", 120),
                ("end", 120),
            ):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(margin))
                node.set(qn("w:type"), "dxa")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title = doc.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 78, 121)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    heading_specs = {
        "Heading 1": (15, RGBColor(31, 78, 121), 12, 5),
        "Heading 2": (12.5, RGBColor(55, 55, 55), 10, 4),
        "Heading 3": (11, RGBColor(80, 80, 80), 8, 3),
        "Heading 4": (10.5, RGBColor(95, 95, 95), 6, 2),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def project_metadata(payload: dict[str, Any]) -> list[tuple[str, str]]:
    project = payload.get("project", {})
    return [
        ("English title", str(project.get("title_en", ""))),
        ("Chinese title", str(project.get("title_zh", ""))),
        ("Manuscript type", str(project.get("manuscript_type", ""))),
        ("Target journal", str(project.get("target_journal", ""))),
        ("Authors", str(project.get("authors", ""))),
        ("Version", str(project.get("version", ""))),
        ("Date", str(project.get("date", ""))),
    ]


def add_title_block(
    doc: Document, title: str, subtitle: str, payload: dict[str, Any]
) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)
    if subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(subtitle)
        run.bold = True
        run.font.size = Pt(12)

    doc.add_paragraph()
    add_key_value_table(doc, project_metadata(payload))
    doc.add_page_break()


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].paragraphs[0].add_run(value or "")
        set_cell_shading(cells[0], "D9EAF7")
    set_table_geometry(table, [1.6, 5.1])


def add_text(doc: Document, value: Any, bold_label: str | None = None) -> None:
    if value is None:
        return
    values = value if isinstance(value, list) else [value]
    for item in values:
        p = doc.add_paragraph()
        if bold_label:
            p.add_run(f"{bold_label}: ").bold = True
        p.add_run(str(item))


def add_section_list(
    doc: Document, sections: list[dict[str, Any]], level_offset: int = 0
) -> None:
    for section in sections:
        heading = str(section.get("heading", "Untitled section"))
        level = int(section.get("level", 1))
        level = min(max(level + level_offset, 1), 4)
        doc.add_heading(heading, level=level)
        for paragraph in section.get("paragraphs", []):
            add_text(doc, paragraph)


def build_manuscript(payload: dict[str, Any], output_path: Path) -> None:
    doc = Document()
    configure_document(doc)
    project = payload.get("project", {})
    add_title_block(
        doc,
        "Bilingual Plant-Science Manuscript",
        f"{project.get('title_en', '')} / {project.get('title_zh', '')}",
        payload,
    )

    doc.add_heading("English Manuscript", level=1)
    add_section_list(
        doc, payload.get("manuscript", {}).get("english", []), level_offset=1
    )
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Academic Chinese Manuscript", level=1)
    add_section_list(
        doc, payload.get("manuscript", {}).get("chinese", []), level_offset=1
    )
    doc.save(output_path)


def build_review(payload: dict[str, Any], output_path: Path) -> None:
    doc = Document()
    configure_document(doc)
    project = payload.get("project", {})
    add_title_block(
        doc,
        "Peer Review, Author Response, and Re-Review",
        str(project.get("title_en", "")),
        payload,
    )
    review = payload.get("review", {})
    doc.add_heading("Editorial Synthesis", level=1)
    add_text(doc, review.get("editorial_synthesis", ""))

    for reviewer in review.get("reviewers", []):
        doc.add_heading(str(reviewer.get("name", "Reviewer")), level=1)
        add_text(doc, reviewer.get("overall_assessment", ""), "Overall assessment")
        for comment in reviewer.get("comments", []):
            comment_id = str(comment.get("id", "Comment"))
            severity = str(comment.get("severity", ""))
            doc.add_heading(f"{comment_id} - {severity}", level=2)
            fields = [
                ("Preserved reviewer comment", comment.get("comment", "")),
                ("Author response", comment.get("response", "")),
                ("Manuscript change", comment.get("change", "")),
                ("Change location", comment.get("location", "")),
                ("Re-review verification", comment.get("verification", "")),
                ("Status", comment.get("status", "")),
            ]
            for heading, value in fields:
                doc.add_heading(heading, level=3)
                add_text(doc, value)

    doc.add_heading("Final Re-Review Decision", level=1)
    add_text(doc, review.get("final_decision", ""))
    doc.add_heading("Residual Risks", level=1)
    add_text(doc, review.get("residual_risks", []))
    doc.save(output_path)


def add_mapping_table(
    doc: Document,
    rows: list[dict[str, Any]],
    columns: list[tuple[str, str]],
    widths_in: list[float] | None = None,
) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, (_, label) in enumerate(columns):
        cell = header.cells[idx]
        cell.text = label
        set_cell_shading(cell, "D9EAF7")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, (key, _) in enumerate(columns):
            cells[idx].text = str(row.get(key, ""))
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths_in is None:
        widths_in = [6.7 / len(columns)] * len(columns)
    set_table_geometry(table, widths_in)


def build_literature(payload: dict[str, Any], output_path: Path) -> None:
    doc = Document()
    configure_document(doc)
    project = payload.get("project", {})
    add_title_block(
        doc,
        "Literature Audit and Full-Paper Close-Reading Record",
        str(project.get("title_en", "")),
        payload,
    )
    literature = payload.get("literature", {})

    doc.add_heading("Search and Screening Record", level=1)
    add_key_value_table(
        doc,
        [(str(k).replace("_", " ").title(), str(v)) for k, v in literature.get("search_record", {}).items()],
    )
    doc.add_heading("Citation Audit Summary", level=1)
    add_key_value_table(
        doc,
        [(str(k).replace("_", " ").title(), str(v)) for k, v in literature.get("audit_summary", {}).items()],
    )
    doc.add_heading("Claim-Citation Matrix", level=1)
    add_mapping_table(
        doc,
        literature.get("claim_citation_matrix", []),
        [
            ("claim_id", "Claim ID"),
            ("manuscript_claim", "Manuscript Claim"),
            ("citation", "Citation"),
            ("support_grade", "Support"),
            ("status", "Status"),
            ("boundary", "Boundary"),
        ],
        [0.55, 1.55, 0.95, 0.75, 0.65, 2.25],
    )

    for index, paper in enumerate(literature.get("papers", []), start=1):
        citation_key = str(paper.get("citation_key", f"Paper {index}"))
        title = str(paper.get("title", "Untitled paper"))
        doc.add_page_break()
        doc.add_heading(f"Paper {index}: {citation_key} - {title}", level=1)
        doc.add_heading("Metadata", level=2)
        metadata_keys = (
            "authors",
            "journal",
            "year",
            "doi",
            "pmid",
            "url",
            "access_status",
            "publication_status",
            "plant_context",
        )
        add_key_value_table(
            doc,
            [(key.replace("_", " ").title(), str(paper.get(key, ""))) for key in metadata_keys],
        )

        doc.add_heading("Full-Paper Close Reading", level=2)
        for reading in paper.get("close_reading", []):
            doc.add_heading(str(reading.get("section", "Section")), level=3)
            for label, key in (
                ("Source anchors", "source_anchors"),
                ("English analytical note", "note_en"),
                ("Academic Chinese note", "note_zh"),
                ("Key evidence", "key_evidence"),
                ("Figures and tables", "figures_tables"),
                ("Limitations", "limitations"),
                ("Manuscript uses", "manuscript_uses"),
            ):
                add_text(doc, reading.get(key, ""), label)

        doc.add_heading("Citation Uses in the Manuscript", level=2)
        for use_index, use in enumerate(paper.get("citation_uses", []), start=1):
            claim_id = str(use.get("claim_id", f"Use {use_index}"))
            doc.add_heading(f"Citation Use {use_index}: {claim_id}", level=3)
            for label, key in (
                ("Manuscript sentence", "manuscript_sentence"),
                ("Manuscript location", "manuscript_location"),
                ("Source anchor", "source_anchor"),
                ("Original source passage", "original_excerpt"),
                ("Academic Chinese translation", "translation_zh"),
                ("Support grade", "support_grade"),
                ("Evidence boundary", "boundary"),
            ):
                add_text(doc, use.get(key, ""), label)

    doc.add_heading("Unresolved Evidence Gaps", level=1)
    add_text(doc, literature.get("unresolved_evidence_gaps", []))
    doc.save(output_path)


def collect_strings(value: Any) -> list[str]:
    if isinstance(value, dict):
        result: list[str] = []
        for child in value.values():
            result.extend(collect_strings(child))
        return result
    if isinstance(value, list):
        result = []
        for child in value:
            result.extend(collect_strings(child))
        return result
    return [str(value)] if value is not None else []


def validate_source_excerpt_lengths(payload: dict[str, Any]) -> list[str]:
    warnings: list[str] = []
    for paper in payload.get("literature", {}).get("papers", []):
        access = str(paper.get("access_status", "")).lower()
        unrestricted = access in {"user-provided", "open-access", "licensed"}
        if unrestricted:
            continue
        for use in paper.get("citation_uses", []):
            excerpt = str(use.get("original_excerpt", ""))
            word_count = len(re.findall(r"\b[\w'-]+\b", excerpt))
            if word_count > 25:
                warnings.append(
                    f"{paper.get('citation_key', 'paper')} {use.get('claim_id', '')}: "
                    f"original excerpt has {word_count} words without an unrestricted access basis"
                )
    return warnings


def validate_docx(path: Path) -> dict[str, Any]:
    doc = Document(path)
    title_paragraphs = [
        p.text.strip() for p in doc.paragraphs if p.style and p.style.name == "Title"
    ]
    heading_levels: list[int] = []
    heading_errors: list[str] = []
    last_level = 0
    for p in doc.paragraphs:
        match = re.fullmatch(r"Heading ([1-4])", p.style.name if p.style else "")
        if not match:
            continue
        level = int(match.group(1))
        heading_levels.append(level)
        if last_level and level > last_level + 1:
            heading_errors.append(
                f"Heading level jumps from {last_level} to {level}: {p.text[:80]}"
            )
        last_level = level
    return {
        "path": str(path.resolve()),
        "reopens": True,
        "nonempty_title": bool(title_paragraphs and title_paragraphs[0]),
        "heading_count": len(heading_levels),
        "heading_hierarchy_errors": heading_errors,
        "table_count": len(doc.tables),
        "paragraph_count": len(doc.paragraphs),
    }


def build_package(payload: dict[str, Any], output_dir: Path) -> dict[str, Any]:
    output_dir.mkdir(parents=True, exist_ok=True)
    paths = {key: output_dir / name for key, name in OUTPUT_NAMES.items()}
    build_manuscript(payload, paths["manuscript"])
    build_review(payload, paths["review"])
    build_literature(payload, paths["literature"])

    all_text = "\n".join(collect_strings(payload))
    placeholders = sorted(
        {
            match.group(0)
            for pattern in PLACEHOLDER_PATTERNS
            for match in re.finditer(pattern, all_text, flags=re.IGNORECASE)
        }
    )
    report = {
        "generated_on": date.today().isoformat(),
        "files": [validate_docx(path) for path in paths.values()],
        "placeholder_warnings": placeholders,
        "excerpt_warnings": validate_source_excerpt_lengths(payload),
        "render_status": "not-run",
        "release_status": "STRUCTURAL_PASS_WITH_WARNINGS"
        if placeholders
        else "STRUCTURAL_PASS_RENDER_REQUIRED",
    }
    report_path = output_dir / "word_package_validation.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Build the three plant-article-writing Word deliverables."
    )
    parser.add_argument("input_json", nargs="?", type=Path)
    parser.add_argument("--output-dir", type=Path, default=Path("output"))
    parser.add_argument("--init-template", type=Path)
    args = parser.parse_args()

    if args.init_template:
        args.init_template.parent.mkdir(parents=True, exist_ok=True)
        args.init_template.write_text(
            json.dumps(template_payload(), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        print(f"Template written: {args.init_template.resolve()}")
        return 0

    if not args.input_json:
        parser.error("input_json is required unless --init-template is used")

    payload = json.loads(args.input_json.read_text(encoding="utf-8-sig"))
    report = build_package(payload, args.output_dir)
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
